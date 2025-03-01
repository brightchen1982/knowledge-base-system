import os
import logging
from typing import Dict, Any, Tuple, List, Optional, BinaryIO
import tempfile
import docx
import re
from .base import DocumentProcessor, register_processor

logger = logging.getLogger(__name__)

@register_processor(extensions=['.docx', '.doc'])
class DocxProcessor(DocumentProcessor):
    """Word文档处理器"""
    
    def process(self, file: BinaryIO) -> Tuple[str, Dict[str, Any]]:
        """
        处理Word文件，提取文本和元数据
        
        Args:
            file: Word文件对象
            
        Returns:
            Tuple[str, Dict[str, Any]]: 提取的文本和元数据
        """
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            # 写入数据
            tmp.write(file.read())
            file.seek(0)  # 重置文件指针
            tmp_path = tmp.name
        
        try:
            # 打开Word文档
            doc = docx.Document(tmp_path)
            
            # 提取元数据
            core_properties = doc.core_properties
            metadata = {
                "title": core_properties.title if hasattr(core_properties, 'title') else "",
                "author": core_properties.author if hasattr(core_properties, 'author') else "",
                "comments": core_properties.comments if hasattr(core_properties, 'comments') else "",
                "keywords": core_properties.keywords if hasattr(core_properties, 'keywords') else "",
                "subject": core_properties.subject if hasattr(core_properties, 'subject') else "",
                "last_modified_by": core_properties.last_modified_by if hasattr(core_properties, 'last_modified_by') else "",
                "created": str(core_properties.created) if hasattr(core_properties, 'created') else "",
                "modified": str(core_properties.modified) if hasattr(core_properties, 'modified') else "",
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
            }
            
            # 提取文本
            extracted_text = ""
            
            # 提取标题
            if doc.paragraphs and doc.paragraphs[0].style.name.startswith('Heading'):
                metadata["document_title"] = doc.paragraphs[0].text
            
            # 提取段落
            for para in doc.paragraphs:
                # 添加段落文本
                if para.text.strip():
                    extracted_text += para.text + "\n\n"
            
            # 提取表格内容
            table_count = 0
            for table in doc.tables:
                table_count += 1
                extracted_text += f"\n--- 表格 {table_count} ---\n"
                
                for i, row in enumerate(table.rows):
                    if i == 0:
                        # 表头
                        extracted_text += "| "
                        for cell in row.cells:
                            extracted_text += cell.text + " | "
                        extracted_text += "\n"
                        extracted_text += "|" + "---|" * len(row.cells) + "\n"
                    else:
                        # 表体
                        extracted_text += "| "
                        for cell in row.cells:
                            extracted_text += cell.text + " | "
                        extracted_text += "\n"
                
                extracted_text += "\n"
            
            metadata["table_count"] = table_count
            
            # 清理文本
            extracted_text = self._clean_text(extracted_text)
            
            return extracted_text, metadata
            
        except Exception as e:
            logger.error(f"处理Word文档时出错: {str(e)}")
            raise Exception(f"Word文档处理失败: {str(e)}")
        finally:
            # 清理临时文件
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def _clean_text(self, text: str) -> str:
        """清理提取的文本"""
        # 删除多余的空格
        text = re.sub(r'\s+', ' ', text)
        
        # 删除多余的换行符
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 保持段落结构
        text = re.sub(r'(\w) *\n *(\w)', r'\1 \2', text)
        
        return text 
